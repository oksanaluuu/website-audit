"""
report.py — generates a .docx audit report from collected data + audit findings.

Usage (called by Claude Code after analysis):
    python src/report.py <domain> --findings findings.json

Or pass findings inline as a JSON string:
    python src/report.py <domain>
    (will look for output/<domain>/findings.json automatically)
"""

import sys
import json
import os
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

AGENCY_NAME = os.getenv("AGENCY_NAME", "Your Agency Name")

# Brand colors
COLOR_DARK = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_ACCENT = RGBColor(0x16, 0x21, 0x3E)
COLOR_HIGHLIGHT = RGBColor(0xE9, 0x4F, 0x37)
COLOR_LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
COLOR_MID_GRAY = RGBColor(0x88, 0x88, 0x88)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_cover_page(doc, domain: str, data: dict):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("WEBSITE AUDIT REPORT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(domain)
    run2.font.size = Pt(16)
    run2.font.color.rgb = COLOR_HIGHLIGHT

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(datetime.now().strftime("%B %d, %Y")).font.color.rgb = COLOR_MID_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    prepared_p = doc.add_paragraph()
    prepared_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = prepared_p.add_run(f"Prepared by {AGENCY_NAME}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = COLOR_MID_GRAY
    run3.italic = True

    doc.add_page_break()


def add_section_heading(doc, text: str):
    p = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.color.rgb = COLOR_DARK
    run.font.size = Pt(16)
    return p


def add_sub_heading(doc, text: str):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.color.rgb = COLOR_ACCENT
    run.font.size = Pt(13)
    return p


def add_finding(doc, issue: str, impact: str, suggestion: str):
    """Add a single finding block: issue label + body."""
    p = doc.add_paragraph()
    run_label = p.add_run("Issue: ")
    run_label.bold = True
    run_label.font.color.rgb = COLOR_HIGHLIGHT
    p.add_run(issue)

    p2 = doc.add_paragraph()
    run_label2 = p2.add_run("Impact: ")
    run_label2.bold = True
    p2.add_run(impact)

    p3 = doc.add_paragraph()
    run_label3 = p3.add_run("Suggestion: ")
    run_label3.bold = True
    run_label3.font.color.rgb = RGBColor(0x20, 0x80, 0x20)
    p3.add_run(suggestion)

    doc.add_paragraph()


def add_screenshot(doc, path: str, caption: str):
    if path and Path(path).exists():
        try:
            doc.add_picture(path, width=Inches(6.0))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = COLOR_MID_GRAY
            cap.runs[0].italic = True
            doc.add_paragraph()
        except Exception as e:
            doc.add_paragraph(f"[Screenshot: {caption} — could not embed: {e}]")


def add_quick_wins_table(doc, quick_wins: list):
    if not quick_wins:
        doc.add_paragraph("No quick wins identified.")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Priority", "Issue", "Effort", "Expected Impact"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr_cells[i], "1A1A2E")

    for win in quick_wins:
        row_cells = table.add_row().cells
        row_cells[0].text = win.get("priority", "")
        row_cells[1].text = win.get("issue", "")
        row_cells[2].text = win.get("effort", "")
        row_cells[3].text = win.get("impact", "")

    doc.add_paragraph()


def generate_report(domain: str, data: dict, findings: dict, output_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover Page ──────────────────────────────────────────────────────────
    add_cover_page(doc, domain, data)

    # ── 1. Executive Summary ────────────────────────────────────────────────
    add_section_heading(doc, "1. Executive Summary")
    doc.add_paragraph(findings.get("executive_summary", ""))
    add_horizontal_rule(doc)

    # ── 2. First Impressions & Visual Design ────────────────────────────────
    add_section_heading(doc, "2. First Impressions & Visual Design")

    homepage = data.get("pages", {}).get("homepage", {})
    if homepage.get("desktop_screenshot"):
        add_screenshot(doc, homepage["desktop_screenshot"], f"{domain} — Desktop View")

    for f in findings.get("visual_design", []):
        add_finding(doc, f["issue"], f["impact"], f["suggestion"])
    add_horizontal_rule(doc)

    # ── 3. Navigation & User Flow ───────────────────────────────────────────
    add_section_heading(doc, "3. Navigation & User Flow")
    for f in findings.get("navigation", []):
        add_finding(doc, f["issue"], f["impact"], f["suggestion"])
    add_horizontal_rule(doc)

    # ── 4. Product Catalog ──────────────────────────────────────────────────
    add_section_heading(doc, "4. Product Catalog")
    products = data.get("pages", {}).get("products", {})
    if products.get("desktop_screenshot"):
        add_screenshot(doc, products["desktop_screenshot"], f"{domain} — Product Catalog")
    for f in findings.get("product_catalog", []):
        add_finding(doc, f["issue"], f["impact"], f["suggestion"])
    add_horizontal_rule(doc)

    # ── 5. RFQ / Lead Generation Forms ─────────────────────────────────────
    add_section_heading(doc, "5. RFQ / Lead Generation Forms")
    contact = data.get("pages", {}).get("contact", {})
    if contact.get("desktop_screenshot"):
        add_screenshot(doc, contact["desktop_screenshot"], f"{domain} — Contact / RFQ Page")
    for f in findings.get("rfq_forms", []):
        add_finding(doc, f["issue"], f["impact"], f["suggestion"])
    add_horizontal_rule(doc)

    # ── 6. Content & Downloads ──────────────────────────────────────────────
    add_section_heading(doc, "6. Content & Downloads")
    for f in findings.get("content_downloads", []):
        add_finding(doc, f["issue"], f["impact"], f["suggestion"])
    add_horizontal_rule(doc)

    # ── 7. SEO Basics ───────────────────────────────────────────────────────
    add_section_heading(doc, "7. SEO Basics")
    add_sub_heading(doc, "What We Observed")

    seo = findings.get("seo_data_summary", {})
    seo_items = [
        ("Page Title", seo.get("meta_title", "Not found")),
        ("Meta Description", seo.get("meta_description", "Not found")),
        ("H1 Tags", ", ".join(seo.get("h1s", [])) or "None found"),
        ("Images missing alt text", str(seo.get("images_missing_alt_count", "N/A"))),
    ]
    for label, value in seo_items:
        p = doc.add_paragraph(style="List Bullet")
        run_l = p.add_run(f"{label}: ")
        run_l.bold = True
        p.add_run(value)

    doc.add_paragraph()
    for f in findings.get("seo", []):
        add_finding(doc, f["issue"], f["impact"], f["suggestion"])
    add_horizontal_rule(doc)

    # ── 8. Mobile Experience ────────────────────────────────────────────────
    add_section_heading(doc, "8. Mobile Experience")

    ps_mobile = data.get("pagespeed", {}).get("mobile", {})
    ps_desktop = data.get("pagespeed", {}).get("desktop", {})
    mobile_score = ps_mobile.get("performance_score", "N/A")
    desktop_score = ps_desktop.get("performance_score", "N/A")

    add_sub_heading(doc, f"Google PageSpeed Scores — Mobile: {mobile_score}/100  |  Desktop: {desktop_score}/100")

    # Embed the PSI screenshots (actual Google report screenshots)
    if ps_mobile.get("screenshot"):
        add_screenshot(doc, ps_mobile["screenshot"], "PageSpeed Insights — Mobile Results")
    if ps_desktop.get("screenshot"):
        add_screenshot(doc, ps_desktop["screenshot"], "PageSpeed Insights — Desktop Results")

    if not ps_mobile.get("screenshot") and homepage.get("mobile_screenshot"):
        add_screenshot(doc, homepage["mobile_screenshot"], f"{domain} — Mobile View")

    for f in findings.get("mobile", []):
        add_finding(doc, f["issue"], f["impact"], f["suggestion"])
    add_horizontal_rule(doc)

    # ── 9. Competitor Benchmarks ────────────────────────────────────────────
    add_section_heading(doc, "9. Competitor Benchmarks")
    doc.add_paragraph(findings.get("competitor_intro", ""))
    doc.add_paragraph()
    for comp in findings.get("competitors", []):
        add_sub_heading(doc, comp.get("name", "Competitor"))
        doc.add_paragraph(comp.get("observation", ""))
        doc.add_paragraph()
    add_horizontal_rule(doc)

    # ── 10. Quick Wins ──────────────────────────────────────────────────────
    add_section_heading(doc, "10. Quick Wins")
    doc.add_paragraph(
        "The following improvements can be implemented without a full redesign "
        "and are expected to have the highest impact on lead generation:"
    )
    doc.add_paragraph()
    add_quick_wins_table(doc, findings.get("quick_wins", []))
    add_horizontal_rule(doc)

    # ── 11. Next Steps ──────────────────────────────────────────────────────
    add_section_heading(doc, "11. Next Steps")
    doc.add_paragraph(findings.get("next_steps_intro", ""))
    doc.add_paragraph()
    for step in findings.get("next_steps", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(step)

    doc.add_paragraph()
    cta = doc.add_paragraph()
    cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cta = cta.add_run(
        findings.get(
            "call_to_action",
            f"We'd love to walk you through these findings in detail and show "
            f"you what a modernized version of your website could look like. "
            f"Book a free 30-minute call with our team."
        )
    )
    run_cta.bold = True
    run_cta.font.size = Pt(12)
    run_cta.font.color.rgb = COLOR_HIGHLIGHT

    # Save
    doc.save(str(output_path))
    print(f"\nReport saved: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python src/report.py \"Client Name\" [YYYY-MM-DD]")
        print("  Date defaults to today if not provided.")
        sys.exit(1)

    client_name = sys.argv[1]
    date_str = sys.argv[2] if len(sys.argv) > 2 else datetime.now().strftime("%Y-%m-%d")
    base = Path(__file__).parent.parent / "output" / client_name / date_str

    data_path = base / "raw" / "data.json"
    findings_path = base / "findings.json"

    if not data_path.exists():
        print(f"Error: No collected data found at {data_path}")
        print("Run collect.py first.")
        sys.exit(1)

    if not findings_path.exists():
        print(f"Error: No findings file found at {findings_path}")
        print("Run the Claude Code audit analysis first to generate findings.json")
        sys.exit(1)

    data = json.loads(data_path.read_text(encoding="utf-8"))
    findings = json.loads(findings_path.read_text(encoding="utf-8"))

    domain = data.get("domain", client_name.lower().replace(" ", "-"))
    slug = client_name.lower().replace(" ", "-")
    output_path = base / f"audit_{slug}_{date_str}.docx"

    generate_report(domain, data, findings, output_path)


if __name__ == "__main__":
    main()
